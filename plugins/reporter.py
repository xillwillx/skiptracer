from docx import Document
from docx.shared import Inches
import __builtin__ as bi

class ReportGenerator:

 def __init__(self):
  pass

 def newdoc(self):
  bi.document = Document()

 def addtitle(self, title):
  bi.document.add_heading(title, 0)

 def writepara(self, paragraph):
   bi.document.add_paragraph(paragraph)

 def addheader(self, header, level):
  bi.document.add_heading(header, level=level)

 def addquote(self, quote):
  bi.document.add_paragraph(quote, style='Intense Quote')

 def unorderedlist(self, data):
  bi.document.add_paragraph(data, style='List Bullet')

 def unorderedlevel(self, data):
  bi.document.add_paragraph(data, style='List Bullet 2')

 def orderedlist(self, data):
  bi.document.add_paragraph(data, style='List Number')

 def addimg(self, imglocation):
  document.add_picture(imglocation) #, width=Inches(1.25))

 def addtable(self, data, rows, cols):
  table = bi.document.add_table(rows=rows, cols=cols)
  hdr_cells = table.rows[0].cells
  hdr_cells[0].text = 'Qty'
  hdr_cells[1].text = 'Id'
  hdr_cells[2].text = 'Desc'
  for qty, id, desc in records:
   row_cells = table.add_row().cells
   row_cells[0].text = str(qty)
   row_cells[1].text = id
   row_cells[2].text = desc

 def addlinebrk(self):
  bi.document.add_page_break()

 def savefile(self, filename):
  bi.document.save(filename)
